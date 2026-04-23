from __future__ import annotations

import re
from datetime import datetime
from html.parser import HTMLParser
from io import BytesIO
from typing import Any

from docx import Document


def _normalize_text(value: str) -> str:
    return re.sub(r"\s+", " ", value).strip()


def _safe_filename_title(value: str) -> str:
    cleaned = re.sub(r'[<>:"/\\\\|?*]+', "", value).strip()
    cleaned = re.sub(r"\s+", " ", cleaned)
    return cleaned or "article"


def _outline_summary(task: dict[str, Any]) -> str:
    raw = str(task.get("keyword") or "").strip()
    if not raw:
        return ""
    first_line = raw.splitlines()[0].strip()
    compact = _normalize_text(first_line or raw)
    return compact[:160].rstrip() + ("..." if len(compact) > 160 else "")


class _HTMLBlockParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.blocks: list[dict[str, Any]] = []
        self.current_block: dict[str, Any] | None = None
        self.bold_depth = 0

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        if tag in {"h1", "h2", "h3", "p", "li"}:
            self.current_block = {"tag": tag, "segments": []}
            return
        if tag in {"strong", "b"}:
            self.bold_depth += 1

    def handle_endtag(self, tag: str) -> None:
        if tag in {"strong", "b"}:
            self.bold_depth = max(0, self.bold_depth - 1)
            return
        if tag in {"h1", "h2", "h3", "p", "li"} and self.current_block:
            text = "".join(segment["text"] for segment in self.current_block["segments"])
            if _normalize_text(text):
                self.blocks.append(self.current_block)
            self.current_block = None

    def handle_data(self, data: str) -> None:
        if not self.current_block:
            return
        normalized = re.sub(r"\s+", " ", data)
        if not normalized.strip():
            if self.current_block["segments"]:
                self.current_block["segments"].append({"text": " ", "bold": False})
            return
        self.current_block["segments"].append({"text": normalized, "bold": self.bold_depth > 0})


class DocExportService:
    def build_docx(self, task: dict[str, Any]) -> tuple[bytes, str]:
        article = task.get("article") or {}
        title = str(article.get("title") or task.get("keyword") or "Article").strip() or "Article"
        timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
        filename = f"{_safe_filename_title(title)}-{timestamp}.docx"

        document = Document()
        self._add_meta(document, task, article)
        document.add_paragraph("")
        document.add_paragraph(title, style="Heading 1")
        self._add_article_body(document, article)

        buffer = BytesIO()
        document.save(buffer)
        return buffer.getvalue(), filename

    def _add_meta(self, document: Document, task: dict[str, Any], article: dict[str, Any]) -> None:
        rows = [
            ("Title", str(article.get("title") or "").strip()),
            ("Outline Summary", _outline_summary(task)),
            ("Meta Title", str(article.get("meta_title") or "").strip()),
            ("Meta Description", str(article.get("meta_description") or "").strip()),
        ]
        for label, value in rows:
            paragraph = document.add_paragraph()
            label_run = paragraph.add_run(f"{label}: ")
            label_run.bold = True
            paragraph.add_run(value)

    def _add_article_body(self, document: Document, article: dict[str, Any]) -> None:
        parser = _HTMLBlockParser()
        parser.feed(str(article.get("raw_html") or article.get("html") or ""))

        first_h1_skipped = False
        article_title = _normalize_text(str(article.get("title") or ""))
        for block in parser.blocks:
            block_text = _normalize_text("".join(segment["text"] for segment in block["segments"]))
            if not block_text:
                continue
            if block["tag"] == "h1" and not first_h1_skipped and block_text == article_title:
                first_h1_skipped = True
                continue

            if block["tag"] == "h1":
                paragraph = document.add_paragraph(style="Heading 1")
            elif block["tag"] == "h2":
                paragraph = document.add_paragraph(style="Heading 2")
            elif block["tag"] == "h3":
                paragraph = document.add_paragraph(style="Heading 3")
            elif block["tag"] == "li":
                paragraph = document.add_paragraph(style="List Bullet")
            else:
                paragraph = document.add_paragraph(style="Normal")

            for segment in block["segments"]:
                text = segment["text"]
                if not text:
                    continue
                run = paragraph.add_run(text)
                if segment["bold"]:
                    run.bold = True

